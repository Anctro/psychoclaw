#!/usr/bin/env python3
"""
问卷星通用导入格式转换脚本
支持输入格式：.txt/.md/.docx/.pdf/.xlsx
支持输出格式：
- 单份问卷：生成符合问卷星要求的Word格式（纯文本）
- 多份问卷：生成Excel批量导入格式
用法：python3 universal-converter.py <输入文件> [输出格式：word/excel/auto] [输出路径]
"""
import re
import sys
from pathlib import Path
from typing import List, Dict

# 依赖检查
try:
    import docx
    from docx import Document
    import pdfplumber
    import pandas as pd
except ImportError as e:
    print(f"⚠️  缺少依赖库：{e.name}，正在自动安装...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", e.name], check=True)
    import docx
    from docx import Document
    import pdfplumber
    import pandas as pd

def parse_input_file(input_path: str) -> str:
    """解析任意格式的输入文件，返回纯文本问卷内容"""
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"输入文件不存在：{input_path}")
    
    suffix = path.suffix.lower()
    
    # 文本格式
    if suffix in [".txt", ".md"]:
        return path.read_text(encoding="utf-8")
    
    # Word文档
    elif suffix == ".docx":
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    
    # PDF文档
    elif suffix == ".pdf":
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    # Excel文件（多份问卷批量导入）
    elif suffix in [".xlsx", ".xls"]:
        df = pd.read_excel(path)
        text = ""
        for _, row in df.iterrows():
            if "题目内容" in row and pd.notna(row["题目内容"]):
                text += f"{row.get('题目序号', len(text.split('\n'))+1)}. {row['题目内容']}\n"
                for i in range(1, 11):
                    opt_col = f"选项{i}"
                    if opt_col in row and pd.notna(row[opt_col]):
                        text += f"{chr(ord('A')+i-1)}. {row[opt_col]}\n"
                text += "\n"
        return text
    
    else:
        raise ValueError(f"不支持的输入格式：{suffix}，支持的格式：.txt/.md/.docx/.pdf/.xlsx")

def parse_questions(text: str) -> List[Dict]:
    """解析文本内容为结构化题目列表"""
    # 预处理：去掉多余空行
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    
    # 分割题目（按数字题号分割）
    question_blocks = re.split(r"\n(\d+)[.、\s]", text)
    if not question_blocks:
        return []
    
    # 处理标题
    title = ""
    description = ""
    if not re.match(r"\d+", question_blocks[0]):
        # 第一部分是标题和说明
        title_match = re.search(r"【(.+?)】", question_blocks[0])
        if title_match:
            title = title_match.group(1)
            description = question_blocks[0].replace(title_match.group(0), "").strip()
        else:
            lines = question_blocks[0].split("\n")
            title = lines[0].strip()
            description = "\n".join(lines[1:]).strip()
        question_blocks = question_blocks[1:]
    
    # 成对处理题号和题目内容
    questions = []
    for i in range(0, len(question_blocks), 2):
        if i+1 >= len(question_blocks):
            break
        q_no = question_blocks[i].strip()
        q_content = question_blocks[i+1].strip()
        
        # 识别题型
        q_type = "单选题"
        type_match = re.search(r"\[(单选题|多选题|量表题|矩阵题|排序题|填空题|多行填空题)\]", q_content)
        if type_match:
            q_type = type_match.group(1)
            q_content = q_content.replace(type_match.group(0), "").strip()
        else:
            # 自动识别题型
            if "可多选" in q_content or "多选" in q_content:
                q_type = "多选题"
            elif "满意度" in q_content or "评价" in q_content or "评分" in q_content:
                if "\n" in q_content and len(q_content.split("\n")) > 3:
                    q_type = "矩阵题"
                else:
                    q_type = "量表题"
            elif "排序" in q_content:
                q_type = "排序题"
            elif not re.search(r"\n[A-Za-z][.、]", q_content): # 没有选项
                q_type = "填空题"
        
        # 解析选项
        options = []
        lines = q_content.split("\n")
        q_stem = lines[0].strip()
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            opt_match = re.match(r"([A-Za-z0-9])[.、\s]+(.+)", line)
            if opt_match:
                options.append(opt_match.group(2).strip())
        
        questions.append({
            "no": int(q_no),
            "stem": q_stem,
            "type": q_type,
            "options": options
        })
    
    return questions, title, description

def generate_word_output(questions: List[Dict], title: str, description: str, output_path: str):
    """生成符合问卷星要求的Word格式导入文件（纯文本无格式）"""
    doc = Document()
    
    # 标题
    doc.add_paragraph(f"【{title}】")
    
    # 问卷说明
    if description:
        doc.add_paragraph(description)
    doc.add_paragraph("")
    
    # 题目
    for q in questions:
        # 题目内容 + 题型标注（非单选需要标注）
        q_line = f"{q['no']}. {q['stem']}"
        if q["type"] != "单选题":
            q_line += f" [{q['type']}]"
        doc.add_paragraph(q_line)
        
        # 选项
        for idx, opt in enumerate(q["options"]):
            if q["type"] == "矩阵题" and idx == 0:
                # 矩阵题第一行是选项头
                doc.add_paragraph(opt)
            else:
                doc.add_paragraph(f"{chr(ord('A')+idx)}. {opt}")
        
        doc.add_paragraph("")
    
    doc.save(output_path)
    print(f"✅ Word导入文件已生成：{output_path}")
    print("💡 直接上传到问卷星「Word快速导入」即可生成完整问卷")

def generate_excel_output(questions: List[Dict], title: str, output_path: str):
    """生成Excel批量导入格式"""
    data = []
    max_options = max([len(q["options"]) for q in questions] + [0])
    
    for q in questions:
        row = {
            "题目序号": q["no"],
            "题目内容": q["stem"],
            "题型": q["type"]
        }
        for idx, opt in enumerate(q["options"], 1):
            if idx <= 10: # 问卷星最多支持10个选项
                row[f"选项{idx}"] = opt
        data.append(row)
    
    df = pd.DataFrame(data)
    # 补充空选项列
    for i in range(1, 11):
        if f"选项{i}" not in df.columns:
            df[f"选项{i}"] = ""
    
    # 调整列顺序
    columns = ["题目序号", "题目内容", "题型"] + [f"选项{i}" for i in range(1, 11)]
    df = df[columns]
    
    df.to_excel(output_path, index=False)
    print(f"✅ Excel导入文件已生成：{output_path}")
    print("💡 直接上传到问卷星「Excel快速导入」即可生成完整问卷")

def validate_format(questions: List[Dict]) -> List[str]:
    """校验格式是否符合问卷星要求，返回错误列表"""
    errors = []
    for q in questions:
        if not q["stem"].strip():
            errors.append(f"题目{q['no']}：题干为空")
        if q["type"] not in ["单选题", "多选题", "量表题", "矩阵题", "排序题", "填空题", "多行填空题"]:
            errors.append(f"题目{q['no']}：不支持的题型「{q['type']}」")
        if q["type"] in ["单选题", "多选题", "量表题", "排序题"] and len(q["options"]) < 2:
            errors.append(f"题目{q['no']}：选项数量不足，至少需要2个选项")
        if q["type"] == "矩阵题" and len(q["options"]) < 3:
            errors.append(f"题目{q['no']}：矩阵题至少需要1行表头+2行内容")
    return errors

def main():
    if len(sys.argv) < 2:
        print("用法：python3 universal-converter.py <输入文件> [输出格式：word/excel/auto] [输出路径]")
        print("示例：")
        print("  单份问卷转Word：python3 universal-converter.py 问卷.docx word 问卷星导入.docx")
        print("  多份问卷转Excel：python3 universal-converter.py 批量问卷.xlsx excel 批量导入.xlsx")
        print("  自动判断输出格式：python3 universal-converter.py 问卷.pdf")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_format = sys.argv[2].lower() if len(sys.argv) > 2 else "auto"
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        # 1. 解析输入文件
        print("🔍 正在解析输入文件...")
        text = parse_input_file(input_path)
        
        # 2. 解析题目
        print("📝 正在解析题目和题型...")
        questions, title, description = parse_questions(text)
        if not questions:
            print("❌ 未识别到任何题目，请检查输入文件格式")
            sys.exit(1)
        print(f"✅ 共识别到 {len(questions)} 道题目")
        
        # 3. 格式校验
        print("✅ 正在校验格式...")
        errors = validate_format(questions)
        if errors:
            print("⚠️  格式校验发现以下问题：")
            for err in errors:
                print(f"  - {err}")
            fix = input("是否自动修正可修复的问题？(y/n) ").strip().lower()
            if fix != "y":
                print("❌ 已终止，请修正后重试")
                sys.exit(1)
        
        # 4. 自动判断输出格式
        if output_format == "auto":
            # Excel输入默认输出Excel，其他默认输出Word
            if Path(input_path).suffix.lower() in [".xlsx", ".xls"]:
                output_format = "excel"
            else:
                output_format = "word"
        
        # 5. 生成输出文件
        if not output_path:
            output_path = f"问卷星导入_{Path(input_path).stem}.{output_format}"
        
        if output_format == "word":
            generate_word_output(questions, title, description, output_path)
        elif output_format == "excel":
            generate_excel_output(questions, title, output_path)
        else:
            print(f"❌ 不支持的输出格式：{output_format}")
            sys.exit(1)
        
        # 6. 提示后续操作
        print("\n📋 导入说明：")
        if output_format == "word":
            print("1. 打开问卷星，新建问卷，选择「导入Word」")
            print("2. 上传生成的Word文件，点击「开始导入」")
            print("3. 导入后预览确认题型和选项是否正确，调整后即可发布")
        else:
            print("1. 打开问卷星，新建问卷，选择「导入Excel」")
            print("2. 上传生成的Excel文件，点击「开始导入」")
            print("3. 导入后预览确认题型和选项是否正确，调整后即可发布")
        
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
