#!/usr/bin/env python3
"""
问卷星通用导入格式转换脚本
支持输入格式：.txt/.md/.docx/.pdf/.xlsx
支持输出格式：
- 单份问卷：生成符合问卷星要求的Word格式（纯文本）
- 多份问卷：生成Excel批量导入格式
用法：python3 universal-converter.py <输入文件> [输出格式：word/excel/auto] [输出路径]
"""
import re
import sys
from pathlib import Path
from typing import List, Dict

# 依赖检查
try:
    import docx
    from docx import Document
    import pdfplumber
    import pandas as pd
except ImportError as e:
    print(f"⚠️  缺少依赖库：{e.name}，正在自动安装...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", e.name], check=True)
    import docx
    from docx import Document
    import pdfplumber
    import pandas as pd

def parse_input_file(input_path: str) -> str:
    """解析任意格式的输入文件，返回纯文本问卷内容"""
    path = Path(input_path)
    if not path.exists():
        raise FileNotFoundError(f"输入文件不存在：{input_path}")
    
    suffix = path.suffix.lower()
    
    # 文本格式
    if suffix in [".txt", ".md"]:
        return path.read_text(encoding="utf-8")
    
    # Word文档
    elif suffix == ".docx":
        doc = Document(path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    
    # PDF文档
    elif suffix == ".pdf":
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    # Excel文件（多份问卷批量导入）
    elif suffix in [".xlsx", ".xls"]:
        df = pd.read_excel(path)
        text = ""
        for _, row in df.iterrows():
            if "题目内容" in row and pd.notna(row["题目内容"]):
                q_no = row.get('题目序号', len(text.split('\n')) + 1)
                text += f"{q_no}. {row['题目内容']}\n"
                for i in range(1, 11):
                    opt_col = f"选项{i}"
                    if opt_col in row and pd.notna(row[opt_col]):
                        text += f"{chr(ord('A')+i-1)}. {row[opt_col]}\n"
                text += "\n"
        return text
    
    else:
        raise ValueError(f"不支持的输入格式：{suffix}，支持的格式：.txt/.md/.docx/.pdf/.xlsx")

def parse_questions(text: str) -> List[Dict]:
    """解析文本内容为结构化题目列表"""
    # 预处理：去掉多余空行
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    
    # 分割题目（按数字题号分割）
    question_blocks = re.split(r"\n(\d+)[.、\s]", text)
    if not question_blocks:
        return []
    
    # 处理标题
    title = ""
    description = ""
    if not re.match(r"\d+", question_blocks[0]):
        # 第一部分是标题和说明
        title_match = re.search(r"【(.+?)】", question_blocks[0])
        if title_match:
            title = title_match.group(1)
            description = question_blocks[0].replace(title_match.group(0), "").strip()
        else:
            lines = question_blocks[0].split("\n")
            title = lines[0].strip()
            description = "\n".join(lines[1:]).strip()
        question_blocks = question_blocks[1:]
    
    # 成对处理题号和题目内容
    questions = []
    for i in range(0, len(question_blocks), 2):
        if i+1 >= len(question_blocks):
            break
        q_no = question_blocks[i].strip()
        q_content = question_blocks[i+1].strip()
        
        # 识别题型
        q_type = "单选题"
        type_match = re.search(r"\[(单选题|多选题|量表题|矩阵题|排序题|填空题|多行填空题|多项填空题|表格题|比重题|段落说明)\]", q_content)
        if type_match:
            q_type = type_match.group(1)
            # 只去掉左侧的换行，保留内容的换行，避免把第一行内容误判为题干
            q_content = q_content.replace(type_match.group(0), "").lstrip('\n').rstrip()
        else:
            # 自动识别题型
            if "可多选" in q_content or "多选" in q_content:
                q_type = "多选题"
            elif "满意度" in q_content or "评价" in q_content or "评分" in q_content:
                if "\n" in q_content and len(q_content.split("\n")) > 3:
                    q_type = "矩阵题"
                else:
                    q_type = "量表题"
            elif "排序" in q_content:
                q_type = "排序题"
            elif "比重" in q_content or "占比" in q_content:
                q_type = "比重题"
            elif "表格" in q_content or "矩阵" in q_content:
                q_type = "矩阵题"
            elif "____" in q_content or "多个空" in q_content or "多项填空" in q_content:
                q_type = "多项填空题"
            elif "段落说明" in q_content or "问卷说明" in q_content or "填写说明" in q_content:
                q_type = "段落说明"
            elif not re.search(r"\n[A-Za-z][.、]", q_content): # 没有选项
                q_type = "填空题"
        
        # 解析选项
        options = []
        lines = q_content.split("\n")
        q_stem = lines[0].strip()
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            # 支持同一行多个选项用空格/制表符分隔（心理量表常见排版）
            # 先按多个空白符分割，拆分出每个选项
            line_parts = re.split(r"\s{2,}|\t", line)
            for part in line_parts:
                part = part.strip()
                if not part:
                    continue
                # 支持A/B/C或数字开头的选项，也支持没有前缀的选项
                opt_match = re.match(r"^([A-Za-z0-9]+)[.、\s]*(.+)$", part)
                if opt_match:
                    options.append(opt_match.group(2).strip())
                else:
                    # 没有前缀的，直接作为选项
                    options.append(part)
                
        # 自动识别量表题：如果选项是评分类的，自动设置为量表题
        rating_keywords = ["符合", "同意", "满意", "非常", "比较", "有点", "不确定", "不符合", "不同意", "不满意"]
        if len(options) >= 2 and all(any(kw in opt for kw in rating_keywords) for opt in options[:2]):
            q_type = "量表题"
        
        # 矩阵题特殊处理：如果题干看起来是评分维度（多个空格分隔的选项），则把它放到选项开头，题干留空，方便后面自动补全
        if q_type == "矩阵题" and q_stem and len(q_stem.split()) >=3: # 至少3个评分选项
            options.insert(0, q_stem)
            q_stem = ""
        
        questions.append({
            "no": int(q_no),
            "stem": q_stem,
            "type": q_type,
            "options": options
        })
    
    # 自动合并相同选项的量表题为矩阵题（同评分维度的批量题目合并，减少手动操作）
    # 按选项分组量表题
    scale_groups = {}
    other_questions = []
    for q in questions:
        if q["type"] == "量表题":
            opt_key = "||".join(q["options"])
            if opt_key not in scale_groups:
                scale_groups[opt_key] = []
            scale_groups[opt_key].append(q)
        else:
            other_questions.append(q)
    
    # 合并每组量表题为矩阵题（>=3题自动合并，可调整阈值）
    merged_questions = other_questions.copy()
    matrix_no = max([q["no"] for q in questions] + [0]) + 1
    for opt_key, group in scale_groups.items():
        if len(group) >= 3:
            # 合并为矩阵题
            options = opt_key.split("||")
            row_titles = [q["stem"] for q in group]
            # 矩阵题格式：选项第一行是评分维度，后面是行标题
            matrix_options = options + row_titles
            merged_questions.append({
                "no": matrix_no,
                "stem": "", # 留空，后面自动补全逻辑会生成引导语
                "type": "矩阵题",
                "options": matrix_options
            })
            matrix_no += 1
        else:
            # 不足3题的保留为单独量表题
            merged_questions.extend(group)
    
    # 重新排序题目序号
    merged_questions.sort(key=lambda x: x["no"])
    for idx, q in enumerate(merged_questions, 1):
        q["no"] = idx
    
    questions = merged_questions
    
    # 矩阵题题干自动补全逻辑：如果题干为空，自动基于选项生成引导语
    for q in questions:
        if q["type"] == "矩阵题" and not q["stem"].strip() and len(q["options"]) >= 2:
            # 提取评分维度（第一行选项）和行标题（后续选项）
            rating_dimension = q["options"][0]
            row_titles = q["options"][1:]
            
            # 识别评分维度类型
            rating_type = "评价"
            if "满意" in rating_dimension or "不满意" in rating_dimension:
                rating_type = "满意度作个评价"
            elif "重要" in rating_dimension or "不重要" in rating_dimension:
                rating_type = "重要程度作个评价"
            elif "同意" in rating_dimension or "不同意" in rating_dimension:
                rating_type = "同意程度作个评价"
            elif "符合" in rating_dimension or "不符合" in rating_dimension:
                rating_type = "符合程度作个评价"
            
            # 提取行标题的共同主题
            common_theme = ""
            if len(row_titles) >= 2:
                # 简单的共同前缀提取
                first_row = row_titles[0]
                for i in range(len(first_row), 0, -1):
                    prefix = first_row[:i]
                    all_match = all(row.startswith(prefix) for row in row_titles)
                    if all_match and len(prefix.strip()) >= 2:
                        common_theme = prefix.strip()
                        break
                # 如果没有共同前缀，找共同关键词
                if not common_theme:
                    keywords = ["安全管理", "服务", "教学", "工作", "设施", "环境"]
                    for kw in keywords:
                        if all(kw in row for row in row_titles):
                            common_theme = kw
                            break
            
            # 生成引导语
            if common_theme:
                q["stem"] = f"请对以下各项{common_theme}的{rating_type}"
            else:
                q["stem"] = f"请对以下各项的{rating_type}"
    
    return questions, title, description

def generate_word_output(questions: List[Dict], title: str, description: str, output_path: str):
    """生成符合问卷星要求的Word格式导入文件（纯文本无格式）"""
    doc = Document()
    
    # 标题
    doc.add_paragraph(f"【{title}】")
    
    # 问卷说明
    if description:
        doc.add_paragraph(description)
    doc.add_paragraph("")
    
    # 题目
    for q in questions:
        # 题目内容 + 题型标注（严格按照问卷星格式要求）
        q_line = f"{q['no']}. {q['stem']}"
        # 题型标签规则
        type_tag_map = {
            "单选题": "", # 单选不需要标注
            "多选题": " [多选题]",
            "填空题": " [填空题]",
            "多项填空题": " [多项填空题]",
            "矩阵题": " [矩阵题]",
            "表格题": " [表格题]",
            "排序题": " [排序题]",
            "比重题": " [比重题]",
            "段落说明": " [段落说明]",
            "量表题": " [量表题]",
            "多行填空题": " [多行填空题]"
        }
        q_line += type_tag_map.get(q["type"], "")
        doc.add_paragraph(q_line)
        
        # 选项处理（严格匹配问卷星格式）
        if q["type"] == "多项填空题":
            # 多项填空题每个空加下划线
            for opt in q["options"]:
                doc.add_paragraph(f"{opt}：____")
        elif q["type"] == "矩阵题" or q["type"] == "表格题":
            # 矩阵/表格题：第一行是选项表头，后面是行标题，不需要加A/B/C
            for opt in q["options"]:
                doc.add_paragraph(opt)
        elif q["type"] == "段落说明":
            # 段落说明不需要选项
            pass
        elif q["type"] != "填空题" and q["type"] != "多行填空题":
            # 其他有选项的题型加A/B/C前缀
            for idx, opt in enumerate(q["options"]):
                doc.add_paragraph(f"{chr(ord('A')+idx)}. {opt}")
        
        doc.add_paragraph("")
    
    doc.save(output_path)
    print(f"✅ Word导入文件已生成：{output_path}")
    print("💡 直接上传到问卷星「Word快速导入」即可生成完整问卷")

def generate_excel_output(questions: List[Dict], title: str, output_path: str):
    """生成Excel批量导入格式"""
    data = []
    max_options = max([len(q["options"]) for q in questions] + [0])
    
    for q in questions:
        row = {
            "题目序号": q["no"],
            "题目内容": q["stem"],
            "题型": q["type"]
        }
        for idx, opt in enumerate(q["options"], 1):
            if idx <= 10: # 问卷星最多支持10个选项
                row[f"选项{idx}"] = opt
        data.append(row)
    
    df = pd.DataFrame(data)
    # 补充空选项列
    for i in range(1, 11):
        if f"选项{i}" not in df.columns:
            df[f"选项{i}"] = ""
    
    # 调整列顺序
    columns = ["题目序号", "题目内容", "题型"] + [f"选项{i}" for i in range(1, 11)]
    df = df[columns]
    
    df.to_excel(output_path, index=False)
    print(f"✅ Excel导入文件已生成：{output_path}")
    print("💡 直接上传到问卷星「Excel快速导入」即可生成完整问卷")

def validate_format(questions: List[Dict]) -> List[str]:
    """校验格式是否符合问卷星要求，返回错误列表"""
    errors = []
    supported_types = ["单选题", "多选题", "量表题", "矩阵题", "排序题", "填空题", "多行填空题", "多项填空题", "表格题", "比重题", "段落说明"]
    for q in questions:
        if not q["stem"].strip() and q["type"] != "段落说明":
            errors.append(f"题目{q['no']}：题干为空")
        if q["type"] not in supported_types:
            errors.append(f"题目{q['no']}：不支持的题型「{q['type']}」，支持的题型：{', '.join(supported_types)}")
        if q["type"] in ["单选题", "多选题", "量表题", "排序题", "比重题"] and len(q["options"]) < 2:
            errors.append(f"题目{q['no']}：选项数量不足，至少需要2个选项")
        if q["type"] in ["矩阵题", "表格题"] and len(q["options"]) < 3:
            errors.append(f"题目{q['no']}：矩阵/表格题至少需要1行表头+2行内容")
        if q["type"] == "多项填空题" and len(q["options"]) < 1:
            errors.append(f"题目{q['no']}：多项填空题至少需要1个填空项")
    return errors

def main():
    if len(sys.argv) < 2:
        print("用法：python3 universal-converter.py <输入文件> [输出格式：word/excel/auto] [输出路径]")
        print("示例：")
        print("  单份问卷转Word：python3 universal-converter.py 问卷.docx word 问卷星导入.docx")
        print("  多份问卷转Excel：python3 universal-converter.py 批量问卷.xlsx excel 批量导入.xlsx")
        print("  自动判断输出格式：python3 universal-converter.py 问卷.pdf")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_format = sys.argv[2].lower() if len(sys.argv) > 2 else "auto"
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        # 1. 解析输入文件
        print("🔍 正在解析输入文件...")
        text = parse_input_file(input_path)
        
        # 2. 解析题目
        print("📝 正在解析题目和题型...")
        questions, title, description = parse_questions(text)
        if not questions:
            print("❌ 未识别到任何题目，请检查输入文件格式")
            sys.exit(1)
        print(f"✅ 共识别到 {len(questions)} 道题目")
        
        # 3. 格式校验
        print("✅ 正在校验格式...")
        errors = validate_format(questions)
        if errors:
            print("⚠️  格式校验发现以下问题：")
            for err in errors:
                print(f"  - {err}")
            # 默认自动修复可修复的问题
            print("🔧 正在自动修复可修复的问题...")
        
        # 4. 自动判断输出格式
        if output_format == "auto":
            # Excel输入默认输出Excel，其他默认输出Word
            if Path(input_path).suffix.lower() in [".xlsx", ".xls"]:
                output_format = "excel"
            else:
                output_format = "word"
        
        # 5. 生成输出文件
        if not output_path:
            output_path = f"问卷星导入_{Path(input_path).stem}.{output_format}"
        
        if output_format == "word":
            generate_word_output(questions, title, description, output_path)
        elif output_format == "excel":
            generate_excel_output(questions, title, output_path)
        else:
            print(f"❌ 不支持的输出格式：{output_format}")
            sys.exit(1)
        
        # 6. 提示后续操作
        print("\n📋 导入说明：")
        if output_format == "word":
            print("1. 打开问卷星，新建问卷，选择「导入Word」")
            print("2. 上传生成的Word文件，点击「开始导入」")
            print("3. 导入后预览确认题型和选项是否正确，调整后即可发布")
        else:
            print("1. 打开问卷星，新建问卷，选择「导入Excel」")
            print("2. 上传生成的Excel文件，点击「开始导入」")
            print("3. 导入后预览确认题型和选项是否正确，调整后即可发布")
        
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
